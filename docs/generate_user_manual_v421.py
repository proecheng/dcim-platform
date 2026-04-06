# -*- coding: utf-8 -*-
"""基于 V4.1.0 用户使用说明书生成 V4.2.1 版本 — 使用 OPC XML 方式创建标题"""

import os, sys, copy
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from lxml import etree

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
SRC_DOC = os.path.join(BASE_DIR, 'DCIM系统用户使用说明书_V4.1.0.docx')
OUT_DOC = os.path.join(BASE_DIR, 'DCIM系统用户使用说明书_V4.2.1.docx')
SS_DIR = os.path.join(BASE_DIR, 'screenshots_v421')

doc = Document(SRC_DOC)
body = doc.element.body

# ========== 找到参考样式的段落 XML 作为模板 ==========
h1_template = None
h2_template = None
normal_template = None

for p in doc.paragraphs:
    if p.style and p.style.name == 'Heading 1' and h1_template is None:
        h1_template = copy.deepcopy(p._element)
    if p.style and p.style.name == 'Heading 2' and h2_template is None:
        h2_template = copy.deepcopy(p._element)
    sname = p.style.name if p.style else ''
    if p.text.strip() and sname not in ('Heading 1', 'Heading 2', 'Heading 3') and normal_template is None:
        normal_template = copy.deepcopy(p._element)
    if h1_template is not None and h2_template is not None and normal_template is not None:
        break

print(f'模板: H1={h1_template is not None}, H2={h2_template is not None}, Normal={normal_template is not None}')

def make_heading(text, level):
    """通过复制模板 XML 创建标题"""
    template = h1_template if level == 1 else h2_template
    new_p = copy.deepcopy(template)
    # 清空所有 run
    for r in new_p.findall(qn('w:r')):
        new_p.remove(r)
    # 添加新 run
    r_elem = OxmlElement('w:r')
    # 复制第一个 run 的 rPr
    orig_runs = template.findall(qn('w:r'))
    if orig_runs:
        orig_rpr = orig_runs[0].find(qn('w:rPr'))
        if orig_rpr is not None:
            r_elem.append(copy.deepcopy(orig_rpr))
    t_elem = OxmlElement('w:t')
    t_elem.text = text
    t_elem.set(qn('xml:space'), 'preserve')
    r_elem.append(t_elem)
    new_p.append(r_elem)
    return new_p

def make_para(text):
    """创建普通段落"""
    new_p = copy.deepcopy(normal_template)
    for r in new_p.findall(qn('w:r')):
        new_p.remove(r)
    r_elem = OxmlElement('w:r')
    # 设置字体
    rpr = OxmlElement('w:rPr')
    rfonts = OxmlElement('w:rFonts')
    rfonts.set(qn('w:ascii'), 'Microsoft YaHei')
    rfonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    rfonts.set(qn('w:hAnsi'), 'Microsoft YaHei')
    rpr.append(rfonts)
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), '21')  # 10.5pt
    rpr.append(sz)
    r_elem.append(rpr)
    t_elem = OxmlElement('w:t')
    t_elem.text = text
    t_elem.set(qn('xml:space'), 'preserve')
    r_elem.append(t_elem)
    new_p.append(r_elem)
    return new_p

def make_bullet(text):
    """创建带圆点的列表项"""
    return make_para('• ' + text)

def make_image(image_path, width=Inches(6.2)):
    """插入图片（通过 python-docx API 创建后提取 XML）"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(image_path, width=width)
    elem = p._element
    # 从 body 尾部移走（稍后会插到正确位置）
    body.remove(elem)
    return elem

def make_caption(text):
    """创建图注"""
    p_elem = make_para(text)
    # 设置居中对齐
    ppr = p_elem.find(qn('w:pPr'))
    if ppr is None:
        ppr = OxmlElement('w:pPr')
        p_elem.insert(0, ppr)
    jc = OxmlElement('w:jc')
    jc.set(qn('w:val'), 'center')
    ppr.append(jc)
    # 设置字体颜色为灰色
    for r in p_elem.findall(qn('w:r')):
        rpr = r.find(qn('w:rPr'))
        if rpr is not None:
            color = OxmlElement('w:color')
            color.set(qn('w:val'), '666666')
            rpr.append(color)
            sz = rpr.find(qn('w:sz'))
            if sz is not None:
                sz.set(qn('w:val'), '18')  # 9pt
    return p_elem

# ========== 1. 更新封面 ==========
for para in doc.paragraphs:
    for run in para.runs:
        if 'V4.1.0' in run.text:
            run.text = run.text.replace('V4.1.0', 'V4.2.1')
        if '2026年3月15日' in run.text:
            run.text = run.text.replace('2026年3月15日', '2026年4月6日')

# ========== 2. 更新协议描述 ==========
for para in doc.paragraphs:
    for run in para.runs:
        if 'BACnet/IP' in run.text and 'MS' not in run.text:
            run.text = run.text.replace('BACnet/IP', 'BACnet IP/MS-TP')

# ========== 3. 修订记录 ==========
if doc.tables:
    t = doc.tables[0]
    row = t.add_row()
    c = row.cells
    c[0].text = 'V4.2.1'
    c[1].text = '2026-04-06'
    c[2].text = '新增：多渠道通知管理、BACnet MS/TP网关状态、预测性维护仪表盘'

# ========== 4. 构建新章节元素 ==========
elems = []

# ---- 第10章 多渠道通知管理 ----
elems.append(make_heading('第10章 多渠道通知管理（V4.2.1 新增）', 1))
elems.append(make_para(
    '多渠道通知管理模块提供告警通知的全链路配置与管理能力，'
    '支持邮件、即时通讯（钉钉/企业微信）、短信、语音等多种通知渠道。'
    '导航路径：系统管理 → 通知管理。'))

elems.append(make_heading('10.1 通知策略', 2))
elems.append(make_para(
    '在「通知策略」标签页中，管理员可以创建和管理通知分发策略。'
    '每条策略可配置：告警级别筛选、通知渠道选择、生效时段、接收人范围等。'))
elems.append(make_image(os.path.join(SS_DIR, 'notification-policies.png')))
elems.append(make_caption('图 10-1 通知策略管理'))
elems.append(make_para('主要功能：'))
elems.append(make_bullet('新建策略：点击右上角「新建策略」按钮，配置策略名称、告警级别、通知渠道和生效时段'))
elems.append(make_bullet('启用/禁用：通过状态开关控制策略是否生效'))
elems.append(make_bullet('时段冲突检测：系统自动检测不同策略间的时段冲突（含跨午夜场景）'))
elems.append(make_bullet('站点权限隔离：策略与站点绑定，不同站点独立管理'))

elems.append(make_heading('10.2 通知记录', 2))
elems.append(make_para(
    '在「通知记录」标签页中，可查看所有通知的发送历史。'
    '支持按渠道类型、状态、告警级别、时间范围筛选。'
    '表格列包含：ID、告警级别、渠道、接收方、内容摘要、状态、错误信息、创建时间。'))
elems.append(make_image(os.path.join(SS_DIR, 'notification-records.png')))
elems.append(make_caption('图 10-2 通知记录查询'))

elems.append(make_heading('10.3 渠道状态', 2))
elems.append(make_para(
    '在「渠道状态」标签页中，展示各通知渠道的健康状态和连通性。'
    '支持对单个渠道发送测试消息，验证通知链路是否正常。'))
elems.append(make_image(os.path.join(SS_DIR, 'notification-channels.png')))
elems.append(make_caption('图 10-3 渠道状态'))

elems.append(make_heading('10.4 用户通知联系方式', 2))
elems.append(make_para(
    '在用户管理页面编辑用户时，新增「通知联系方式」标签页，'
    '可为每个用户配置多种通知渠道的联系方式（邮箱、手机号、钉钉ID、企业微信ID等）。'
    '支持从用户个人资料一键导入。'))

elems.append(make_heading('10.5 告警风暴抑制', 2))
elems.append(make_para(
    '当短时间内产生大量同类告警时，系统自动启动风暴抑制机制，'
    '将多条告警合并为一条摘要通知，避免通知轰炸。'))

elems.append(make_heading('10.6 通知渠道升级', 2))
elems.append(make_para(
    '当严重告警在指定时间内未被确认处理，系统自动将通知升级到更高优先级的渠道。'
    '例如：邮件未响应 → 短信通知 → 电话通知。升级链可在通知策略中配置。'))

# ---- 第11章 预测性维护 ----
elems.append(make_heading('第11章 预测性维护仪表盘（V4.2.1 新增）', 1))
elems.append(make_para(
    '预测性维护仪表盘提供设备劣化趋势分析和维护建议能力，'
    '覆盖 HVAC（空调）、UPS、PDU、Battery 四类核心设备。'
    '导航路径：运维管理 → 预测性维护。'))

elems.append(make_heading('11.1 仪表盘总览', 2))
elems.append(make_para(
    '页面顶部展示四个统计卡片：总设备数、健康设备数、关注设备数、预警+危险设备数。'
    '筛选栏支持按设备类型和健康等级过滤。'))
elems.append(make_image(os.path.join(SS_DIR, 'predictive-maintenance.png')))
elems.append(make_caption('图 11-1 预测性维护仪表盘'))

elems.append(make_heading('11.2 设备健康评分', 2))
elems.append(make_para('系统为每台设备计算健康度评分（0-100分），根据分数划分四个等级：'))
elems.append(make_bullet('健康（80-100分）：设备运行正常，无需特别关注'))
elems.append(make_bullet('关注（60-79分）：设备出现轻微劣化，建议关注趋势'))
elems.append(make_bullet('预警（40-59分）：设备劣化明显，建议安排维护'))
elems.append(make_bullet('危险（0-39分）：设备严重劣化，需立即处理'))

elems.append(make_heading('11.3 劣化分析插件', 2))
elems.append(make_para('系统内置四类设备的劣化分析插件：'))
elems.append(make_bullet('HVAC 插件：分析制冷效率下降、能耗异常、温度控制偏差等'))
elems.append(make_bullet('UPS 插件：分析电池容量衰减、转换效率、负载能力等'))
elems.append(make_bullet('PDU 插件：分析分路功率偏差、温升异常、负载不均衡等'))
elems.append(make_bullet('Battery 插件：分析 SOH 衰减趋势、内阻变化、充放电效率等'))

elems.append(make_heading('11.4 维护建议', 2))
elems.append(make_para('系统根据劣化分析结果自动生成维护建议。在设备详情弹窗中可查看：'))
elems.append(make_bullet('健康度评分及趋势'))
elems.append(make_bullet('劣化评分各因子权重分析'))
elems.append(make_bullet('维护建议列表（含紧急程度标记）'))
elems.append(make_bullet('「确认」操作：接受建议并自动创建维护工单'))
elems.append(make_bullet('「拒绝」操作：标记为误报并提供反馈，用于优化模型'))
elems.append(make_para(
    '数据充分度标记：系统会标注数据充分度（完整/中等/有限），'
    '提醒运维人员注意评估结果的可信度。'))

# ---- 第12章 数据源管理增强 ----
elems.append(make_heading('第12章 数据源管理增强（V4.2.1 更新）', 1))
elems.append(make_para(
    '数据源管理（采集配置 → 数据源管理）新增 BACnet MS/TP 协议支持和双层故障隔离功能。'))

elems.append(make_heading('12.1 BACnet MS/TP 网关接入', 2))
elems.append(make_para(
    '支持通过 BACnet MS/TP → IP 协议转换网关接入大金空调等设备。'
    '在数据源管理中新增 BACnet MS/TP 协议类型，支持网关模板和批量创建端点。'))

elems.append(make_heading('12.2 双层故障隔离', 2))
elems.append(make_para('数据源状态新增两种故障状态，实现网关级和设备级的双层故障隔离：'))
elems.append(make_bullet('网关离线（gateway_offline）：网关设备本身失去连接，该网关下所有设备自动级联为离线状态'))
elems.append(make_bullet('设备离线（device_offline）：单台设备通信中断，不影响同网关的其他设备'))
elems.append(make_image(os.path.join(SS_DIR, 'datasource-management.png')))
elems.append(make_caption('图 12-1 数据源管理（含新增状态筛选）'))
elems.append(make_para(
    '状态筛选器现支持 6 种状态：已连接、未连接、通信中断、网关离线、设备离线。'
    '当网关恢复在线时，其下属设备的告警自动级联关闭。'))

# ========== 5. 按正序插入到附录前 ==========
def find_para_idx(text_contains):
    for i, p in enumerate(doc.paragraphs):
        if text_contains in p.text:
            return i
    return -1

appendix_idx = find_para_idx('附录A')
ref = doc.paragraphs[appendix_idx]._element

# 正序插入：每个元素都放在 ref 前面，但要保持顺序
# 方法：先插第一个到 ref 前，然后后续每个插到前一个后面
prev = None
for elem in elems:
    if prev is None:
        ref.addprevious(elem)
    else:
        prev.addnext(elem)
    prev = elem

# ========== 6. 保存 ==========
doc.save(OUT_DOC)
print(f'文档已生成: {OUT_DOC}')
print(f'段落数: {len(doc.paragraphs)}')
print(f'图片数: {len(doc.inline_shapes)}')
